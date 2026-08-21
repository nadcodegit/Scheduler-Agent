from __future__ import annotations

from datetime import date
from pathlib import Path

import docx

from scheduler_agents.models.state import TimesheetData

_MONTHS = {
    name: i
    for i, name in enumerate(
        [
            "January",
            "February",
            "March",
            "April",
            "May",
            "June",
            "July",
            "August",
            "September",
            "October",
            "November",
            "December",
        ],
        start=1,
    )
}


def _period_to_yyyymm(period: str) -> str:
    month_name, year = period.rsplit(" ", 1)
    return f"{year}{_MONTHS[month_name.capitalize()]:02d}"


def _set_paragraph_text(paragraph, text: str) -> None:
    """Overwrite a paragraph's text on its first run, clearing the rest.

    Word splits a single visible line across several runs (spell-check
    markers, revision ids), so writing to run[0].text alone would leave
    stale fragments from the other runs appended after it. Setting run[0]
    and blanking the remainder keeps the first run's formatting (bold,
    color) intact without introducing extra empty runs.
    """

    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def fill_invoice_template(
    template_path: Path,
    output_path: Path,
    data: TimesheetData,
    *,
    vendor_id: str,
    invoice_date: date,
) -> Path:
    """Fill only the known monthly-varying placeholders in the fixed invoice
    template: invoice number, invoice date, job id, and amount/total. Every
    other field (name, address, bank details, bill-to) is the interpreter's
    own static info and is left untouched.

    Cell positions are specific to this template's layout (inspected once,
    not re-derived from label text each run) -- this is a "known fixed
    template" tool, not a generic docx editor.
    """

    doc = docx.Document(str(template_path))

    invoice_number = f"{vendor_id}/{_period_to_yyyymm(data.period)}"
    date_str = invoice_date.strftime("%d/%m/%Y")
    amount_str = f"€ {data.total_amount:.2f}"

    header_table = doc.tables[0]
    _set_paragraph_text(header_table.rows[3].cells[1].paragraphs[1], invoice_number)
    _set_paragraph_text(header_table.rows[3].cells[2].paragraphs[1], date_str)

    line_items_table = doc.tables[2]
    _set_paragraph_text(line_items_table.rows[2].cells[0].paragraphs[0], data.job_id)
    _set_paragraph_text(line_items_table.rows[2].cells[2].paragraphs[0], amount_str)
    _set_paragraph_text(line_items_table.rows[3].cells[2].paragraphs[0], amount_str)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
